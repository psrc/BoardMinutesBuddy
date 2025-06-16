#!/usr/bin/env python3
"""
Board Minutes Buddy - GUI Application for Generating Meeting Minutes
Converts R script functionality to Python with tkinter GUI
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json
import os
import re
import random
from pathlib import Path
import openai
from docx import Document
from docx.shared import Inches
import threading
import configparser

class BoardMinutesBuddy:
    def __init__(self, root):
        self.root = root
        self.root.title("Board Minutes Buddy")
        self.root.geometry("600x500")
        self.root.resizable(True, True)
        
        # Configuration file for persistent settings
        self.config_file = "board_minutes_config.ini"
        self.config = configparser.ConfigParser()
        self.load_settings()
        
        # Variables for file paths
        self.agenda_path = tk.StringVar()
        self.transcript_path = tk.StringVar()
        self.output_dir = tk.StringVar(value=str(Path.home()))
        
        self.setup_gui()
        
    def load_settings(self):
        """Load settings from config file"""
        if os.path.exists(self.config_file):
            self.config.read(self.config_file)
        
        # Set defaults if not present
        if not self.config.has_section('Settings'):
            self.config.add_section('Settings')
        
        if not self.config.has_option('Settings', 'model'):
            self.config.set('Settings', 'model', 'gpt-4.1')
        
        if not self.config.has_option('Settings', 'api_key'):
            self.config.set('Settings', 'api_key', '')
        
        if not self.config.has_option('Settings', 'template_path'):
            self.config.set('Settings', 'template_path', '')
        
        if not self.config.has_option('Settings', 'examples_path'):
            self.config.set('Settings', 'examples_path', '')
    
    def save_settings(self):
        """Save settings to config file"""
        with open(self.config_file, 'w') as f:
            self.config.write(f)
    
    def setup_gui(self):
        """Create the main GUI interface"""
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text="Board Minutes Buddy", 
                               font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # File selection section
        file_frame = ttk.LabelFrame(main_frame, text="File Selection", padding="10")
        file_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        file_frame.columnconfigure(1, weight=1)
        
        # Agenda file
        ttk.Label(file_frame, text="Agenda File:").grid(row=0, column=0, sticky=tk.W, pady=2)
        agenda_entry = ttk.Entry(file_frame, textvariable=self.agenda_path)
        agenda_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=2)
        ttk.Button(file_frame, text="Browse", 
                  command=self.browse_agenda).grid(row=0, column=2, pady=2)
        
        # Transcript file
        ttk.Label(file_frame, text="Transcript File:").grid(row=1, column=0, sticky=tk.W, pady=2)
        transcript_entry = ttk.Entry(file_frame, textvariable=self.transcript_path)
        transcript_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=2)
        ttk.Button(file_frame, text="Browse", 
                  command=self.browse_transcript).grid(row=1, column=2, pady=2)
        
        # Output directory
        ttk.Label(file_frame, text="Output Directory:").grid(row=2, column=0, sticky=tk.W, pady=2)
        output_entry = ttk.Entry(file_frame, textvariable=self.output_dir)
        output_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=2)
        ttk.Button(file_frame, text="Browse", 
                  command=self.browse_output).grid(row=2, column=2, pady=2)
        
        # Settings section
        settings_frame = ttk.LabelFrame(main_frame, text="Settings", padding="10")
        settings_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        settings_frame.columnconfigure(1, weight=1)
        
        # Model selection
        ttk.Label(settings_frame, text="Model:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.model_var = tk.StringVar(value=self.config.get('Settings', 'model'))
        model_entry = ttk.Entry(settings_frame, textvariable=self.model_var)
        model_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=2)
        
        # API Key
        ttk.Label(settings_frame, text="OpenAI API Key:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.api_key_var = tk.StringVar(value=self.config.get('Settings', 'api_key'))
        api_key_entry = ttk.Entry(settings_frame, textvariable=self.api_key_var, show="*")
        api_key_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=2)
        
        # Template Path
        ttk.Label(settings_frame, text="Template File:").grid(row=2, column=0, sticky=tk.W, pady=2)
        self.template_path_var = tk.StringVar(value=self.config.get('Settings', 'template_path'))
        template_entry = ttk.Entry(settings_frame, textvariable=self.template_path_var)
        template_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=2)
        ttk.Button(settings_frame, text="Browse", 
                  command=self.browse_template).grid(row=2, column=2, pady=2)
        
        # Examples Path
        ttk.Label(settings_frame, text="Examples File:").grid(row=3, column=0, sticky=tk.W, pady=2)
        self.examples_path_var = tk.StringVar(value=self.config.get('Settings', 'examples_path'))
        examples_entry = ttk.Entry(settings_frame, textvariable=self.examples_path_var)
        examples_entry.grid(row=3, column=1, sticky=(tk.W, tk.E), padx=(10, 5), pady=2)
        ttk.Button(settings_frame, text="Browse", 
                  command=self.browse_examples).grid(row=3, column=2, pady=2)
        
        # Control buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=3, column=0, columnspan=3, pady=10)
        
        self.generate_button = ttk.Button(button_frame, text="Generate Minutes", 
                                         command=self.generate_minutes_thread)
        self.generate_button.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(button_frame, text="Save Settings", 
                  command=self.save_settings_gui).pack(side=tk.LEFT, padx=5)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        # Status/Log area
        log_frame = ttk.LabelFrame(main_frame, text="Status", padding="5")
        log_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(10, 0))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        main_frame.rowconfigure(5, weight=1)
        
        self.log_text = tk.Text(log_frame, height=8, wrap=tk.WORD)
        scrollbar = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
    def browse_agenda(self):
        """Browse for agenda file"""
        filename = filedialog.askopenfilename(
            title="Select Agenda File",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if filename:
            self.agenda_path.set(filename)
    
    def browse_transcript(self):
        """Browse for transcript file"""
        filename = filedialog.askopenfilename(
            title="Select Transcript File",
            filetypes=[("VTT Files", "*.vtt"), ("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        if filename:
            self.transcript_path.set(filename)
    
    def browse_output(self):
        """Browse for output directory"""
        dirname = filedialog.askdirectory(title="Select Output Directory")
        if dirname:
            self.output_dir.set(dirname)
    
    def browse_template(self):
        """Browse for template file"""
        filename = filedialog.askopenfilename(
            title="Select Word Template File",
            filetypes=[("Word Templates", "*.dotx"), ("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if filename:
            self.template_path_var.set(filename)
    
    def browse_examples(self):
        """Browse for examples JSONL file"""
        filename = filedialog.askopenfilename(
            title="Select Examples JSONL File",
            filetypes=[("JSONL Files", "*.jsonl"), ("JSON Files", "*.json"), ("All Files", "*.*")]
        )
        if filename:
            self.examples_path_var.set(filename)
    
    def save_settings_gui(self):
        """Save settings from GUI"""
        self.config.set('Settings', 'model', self.model_var.get())
        self.config.set('Settings', 'api_key', self.api_key_var.get())
        self.config.set('Settings', 'template_path', self.template_path_var.get())
        self.config.set('Settings', 'examples_path', self.examples_path_var.get())
        self.save_settings()
        self.log("Settings saved successfully!")
    
    def log(self, message):
        """Add message to log area"""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def extract_word_text(self, filepath):
        """Extract text from Word document with improved formatting preservation"""
        try:
            doc = Document(filepath)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Preserve basic structure and detect headings
                    text = paragraph.text.strip()
                    
                    # Check if paragraph style suggests it's a heading
                    if paragraph.style.name.lower().startswith('heading'):
                        # Add markdown-style heading markers based on level
                        level = 1  # default
                        if 'heading 1' in paragraph.style.name.lower():
                            level = 1
                        elif 'heading 2' in paragraph.style.name.lower():
                            level = 2
                        elif 'heading 3' in paragraph.style.name.lower():
                            level = 3
                        text = '#' * level + ' ' + text
                    
                    text_content.append(text)
            
            combined_text = '\n'.join(text_content)
            # Only replace multiple spaces/tabs, but preserve newlines
            combined_text = re.sub(r'[ \t]+', ' ', combined_text)
            # Clean up excessive newlines (3+ becomes 2)
            combined_text = re.sub(r'\n{3,}', '\n\n', combined_text)
            return combined_text.strip()
            
        except Exception as e:
            self.log(f"Error reading Word document {filepath}: {str(e)}")
            return ""
    
    def extract_vtt_text(self, filepath):
        """Extract text from VTT transcript file with improved cleaning"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                lines = file.readlines()
            
            # Remove VTT header and timestamp lines
            # VTT format: WEBVTT, timestamps (00:00:00.000 --> 00:00:00.000), then speaker: text
            content_lines = []
            for line in lines:
                line = line.strip()
                # Skip VTT headers, timestamps, and empty lines
                if not re.match(r'^WEBVTT|^NOTE|^\d{2}:\d{2}:\d{2}\.\d{3}', line) and line:
                    content_lines.append(line)
            
            # Combine speaker attributions and dialogue
            transcript_text = ' '.join(content_lines)
            # Clean up excessive whitespace
            transcript_text = re.sub(r'\s+', ' ', transcript_text)
            return transcript_text.strip()
            
        except Exception as e:
            self.log(f"Error reading VTT file {filepath}: {str(e)}")
            return ""
    
    def load_examples_from_jsonl(self, jsonl_path):
        """Load examples from JSONL file"""
        if not jsonl_path or not os.path.exists(jsonl_path):
            return []
        
        try:
            examples = []
            with open(jsonl_path, 'r', encoding='utf-8') as file:
                for line_num, line in enumerate(file, 1):
                    line = line.strip()
                    if not line:
                        continue
                    
                    try:
                        example = json.loads(line)
                        
                        # Validate structure
                        if "messages" not in example or len(example["messages"]) < 3:
                            self.log(f"Warning: Invalid example structure at line {line_num}")
                            continue
                        
                        examples.append(example)
                        
                    except json.JSONDecodeError as e:
                        self.log(f"Warning: JSON decode error at line {line_num}: {str(e)}")
                        continue
            
            self.log(f"Loaded {len(examples)} examples from JSONL file")
            return examples
            
        except Exception as e:
            self.log(f"Error loading JSONL file {jsonl_path}: {str(e)}")
            return []
    
    def get_examples_file(self):
        """Get the examples file path from settings or find automatically"""
        # First check if user has specified an examples path
        examples_path = self.examples_path_var.get().strip()
        
        if examples_path and os.path.exists(examples_path):
            return examples_path
        
        # If no valid path in settings, try to find automatically
        return self.find_examples_file()
        
    def find_examples_file(self):
        """Find the example.jsonl file automatically"""
        # Look for examples file in common locations
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        search_locations = [
            os.path.join(script_dir, "example.jsonl"),
            os.path.join(os.getcwd(), "example.jsonl"),
            os.path.join(str(Path.home()), "example.jsonl"),
            os.path.join(str(Path.home()), "Documents", "example.jsonl")
        ]
        
        for path in search_locations:
            if os.path.exists(path):
                return path
        
        return None
    
    def get_template_file(self):
        """Get the template file path from settings or find automatically"""
        # First check if user has specified a template path
        template_path = self.template_path_var.get().strip()
        
        if template_path and os.path.exists(template_path):
            return template_path
        
        # If no valid path in settings, try to find automatically
        return self.find_template_file()
        
    def find_template_file(self):
        """Find the minutes_msword.dotx template file automatically"""
        # Look for template in common locations
        script_dir = os.path.dirname(os.path.abspath(__file__))
        
        search_locations = [
            os.path.join(script_dir, "minutes_msword.dotx"),
            os.path.join(script_dir, "minutes_msword.docx"),
            os.path.join(os.getcwd(), "minutes_msword.dotx"),
            os.path.join(os.getcwd(), "minutes_msword.docx"),
            os.path.join(str(Path.home()), "minutes_msword.dotx"),
            os.path.join(str(Path.home()), "minutes_msword.docx"),
            os.path.join(str(Path.home()), "Documents", "minutes_msword.dotx"),
            os.path.join(str(Path.home()), "Documents", "minutes_msword.docx")
        ]
        
        for path in search_locations:
            if os.path.exists(path):
                return path
        
        return None
    
    def generate_minutes_chatgpt(self, agenda_text, transcript_text):
        """Generate minutes using OpenAI API with one-shot learning from JSONL examples"""
        try:
            # Set up OpenAI client
            openai.api_key = self.api_key_var.get()
            
            # Load examples from JSONL file
            examples_file = self.get_examples_file()
            if examples_file:
                examples = self.load_examples_from_jsonl(examples_file)
                self.log(f"Using examples from: {examples_file}")
            else:
                examples = []
                self.log("No examples file found. Using zero-shot approach.")
            
            # System message
            system_msg = {
                "role": "system",
                "content": ("You are an expert at creating structured meeting minutes for public agency board meetings. "
                           "Given an agenda and meeting transcript, you produce concise, professional minutes that "
                           "follow the agenda structure, cogently summarize key discussions without quotes, and "
                           "particularly highlight board actions reached during the meeting ,"
                           "(i.e., motions that are seconded and pass by vote using Robert's Rules of Order). "
                           "Closely follow the style, formatting and conventions you observe in the included example.")
            }
            
            # Build messages for one-shot learning
            messages = [system_msg]
            
            # Add single example if available (one-shot approach)
            if examples and len(examples) > 0:
                # Use first example or random selection
                selected_example = examples[0]  # or random.choice(examples)
                
                if 'messages' in selected_example and len(selected_example['messages']) >= 3:
                    # Add the example user message and assistant response
                    messages.append(selected_example['messages'][1])  # user message
                    messages.append(selected_example['messages'][2])  # assistant message
                    self.log("Using one-shot learning with example")
            
            # Add the current task
            final_task = {
                "role": "user",
                "content": f"Please create meeting minutes based on this agenda and transcript.\n\nAGENDA:\n{agenda_text}\n\nTRANSCRIPT:\n{transcript_text}"
            }
            messages.append(final_task)
            
            self.log("Sending request to OpenAI...")
            
            # Make API call
            response = openai.ChatCompletion.create(
                model=self.model_var.get(),
                messages=messages,
                max_tokens=6500,
                temperature=0.3
            )
            
            minutes_rawtext = response.choices[0].message.content
            self.log("Minutes generated successfully!")
            return minutes_rawtext
            
        except Exception as e:
            self.log(f"Error generating minutes: {str(e)}")
            return None
    
    def minutes_to_word_template(self, minutes_text, output_path):
        """Convert minutes text to formatted Word document using template"""
        try:
            # Get the template file
            template_path = self.get_template_file()
            
            if template_path:
                self.log(f"Using template: {template_path}")
                
                # Check if it's a .dotx template file
                if template_path.lower().endswith('.dotx'):
                    self.log("Warning: .dotx templates are not fully supported by python-docx.")
                    self.log("For best results, please save your template as a .docx file.")
                    self.log("Creating document with basic template-inspired formatting...")
                    
                    # Create new document and apply PSRC-style formatting
                    doc = Document()
                    self.setup_psrc_styles(doc)
                    
                elif template_path.lower().endswith('.docx'):
                    # Load the .docx template
                    doc = Document(template_path)
                    self.log("Loaded .docx template successfully")
                    
                    # Don't clear content - preserve logo and existing formatting
                    self.log("Preserving template logo and styles")
                else:
                    # Unknown template type
                    self.log("Unknown template file type. Creating basic document.")
                    doc = Document()
                    self.setup_basic_styles(doc)
            else:
                self.log("No template file found. Creating document with standard formatting.")
                doc = Document()
                self.setup_basic_styles(doc)
            
            # Parse and add the minutes content
            self.parse_and_add_content(doc, minutes_text)
            
            # Save document
            doc.save(output_path)
            self.log(f"Document saved to: {output_path}")
            
        except Exception as e:
            self.log(f"Error creating Word document: {str(e)}")
            # Fallback: create basic document
            try:
                self.log("Attempting fallback document creation...")
                doc = Document()
                self.setup_basic_styles(doc)
                # Add the raw minutes text as paragraphs
                paragraphs = minutes_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                doc.save(output_path)
                self.log(f"Fallback document saved to: {output_path}")
            except Exception as fallback_error:
                self.log(f"Fallback also failed: {str(fallback_error)}")
    
    def setup_psrc_styles(self, doc):
        """Set up PSRC-style formatting for documents"""
        try:
            # Add PSRC header/logo placeholder
            header_para = doc.add_paragraph()
            header_run = header_para.add_run("PUGET SOUND REGIONAL COUNCIL")
            header_run.bold = True
            header_para.alignment = 1  # Center alignment
            
            # Add some spacing
            doc.add_paragraph()
            
        except Exception as e:
            self.log(f"Error setting up PSRC styles: {str(e)}")
    
    def setup_basic_styles(self, doc):
        """Set up basic document styles"""
        try:
            # Just add the basic title
            doc.add_heading('Meeting Minutes', 0)
        except Exception as e:
            self.log(f"Error setting up basic styles: {str(e)}")
    
    def clear_template_content(self, doc):
        """Clear existing content from template while preserving styles"""
        try:
            # Remove all paragraphs except the first one (which may contain important styles)
            paragraphs_to_remove = []
            for i, paragraph in enumerate(doc.paragraphs):
                if i > 0:  # Keep the first paragraph for style reference
                    paragraphs_to_remove.append(paragraph)
            
            for paragraph in paragraphs_to_remove:
                p = paragraph._element
                p.getparent().remove(p)
                
            self.log("Template content cleared successfully")
            
        except Exception as e:
            self.log(f"Error clearing template content: {str(e)}")
            # If we can't clear, just proceed - the content will be appended
    
    def parse_and_add_content(self, doc, minutes_text):
        """Parse minutes text and add to document with appropriate formatting"""
        try:
            # Split content into lines
            lines = minutes_text.split('\n')
            current_paragraph = []
            in_list = False
            
            for line in lines:
                original_line = line
                line = line.strip()
                
                if not line:
                    # Empty line - end current paragraph if it exists
                    if current_paragraph:
                        paragraph_text = ' '.join(current_paragraph)
                        self.add_formatted_paragraph(doc, paragraph_text, in_list)
                        current_paragraph = []
                        in_list = False
                    continue
                
                # Check for markdown headings first
                if line.startswith('#'):
                    # Add any accumulated paragraph first
                    if current_paragraph:
                        paragraph_text = ' '.join(current_paragraph)
                        self.add_formatted_paragraph(doc, paragraph_text, in_list)
                        current_paragraph = []
                        in_list = False
                    
                    # Process markdown headings
                    heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
                    if heading_match:
                        level = len(heading_match.group(1))
                        heading_text = heading_match.group(2)
                        doc.add_heading(heading_text, level=min(level, 3))
                    continue
                
                # Check for bullet points or list items
                if re.match(r'^[-*+]\s+', line) or re.match(r'^\d+\.\s+', line):
                    # Add any accumulated paragraph first
                    if current_paragraph and not in_list:
                        paragraph_text = ' '.join(current_paragraph)
                        self.add_formatted_paragraph(doc, paragraph_text, False)
                        current_paragraph = []
                    
                    # Process list item
                    list_text = re.sub(r'^[-*+]\s+', '', line)  # Remove bullet
                    list_text = re.sub(r'^\d+\.\s+', '', list_text)  # Remove number
                    self.add_list_item(doc, list_text)
                    in_list = True
                    continue
                
                # Check if line looks like a heading (non-markdown)
                elif self.is_heading(line):
                    # Add any accumulated paragraph first
                    if current_paragraph:
                        paragraph_text = ' '.join(current_paragraph)
                        self.add_formatted_paragraph(doc, paragraph_text, in_list)
                        current_paragraph = []
                        in_list = False
                    
                    # Determine heading level for non-markdown headings
                    heading_level = self.get_heading_level(line)
                    doc.add_heading(line, level=heading_level)
                    continue
                
                else:
                    # Regular content - accumulate into paragraph
                    if in_list and current_paragraph:
                        # If we were in a list but now have regular text, end the list
                        paragraph_text = ' '.join(current_paragraph)
                        self.add_formatted_paragraph(doc, paragraph_text, True)
                        current_paragraph = []
                        in_list = False
                    
                    current_paragraph.append(line)
            
            # Add any remaining paragraph
            if current_paragraph:
                paragraph_text = ' '.join(current_paragraph)
                self.add_formatted_paragraph(doc, paragraph_text, in_list)
                
        except Exception as e:
            self.log(f"Error parsing content: {str(e)}")
            # Fallback: add as single paragraph
            doc.add_paragraph(minutes_text)
    
    def add_list_item(self, doc, text):
        """Add a list item with proper formatting"""
        try:
            # Try to use proper list style
            try:
                p = doc.add_paragraph(style='List Bullet')
            except:
                # If List Bullet style not available, create regular paragraph
                p = doc.add_paragraph()
                p.add_run("• ")  # Add bullet manually
            
            # Process any inline formatting (bold, italic)
            self.add_formatted_text_to_paragraph(p, text)
            
        except Exception as e:
            # Fallback to regular paragraph with bullet
            doc.add_paragraph(f"• {text}")
    
    def add_formatted_paragraph(self, doc, text, is_list_continuation=False):
        """Add a paragraph with formatting support and proper styles"""
        try:
            # Check for special formatting patterns
            if re.match(r'^\*\*Action:\*\*', text, re.IGNORECASE):
                # ACTION items - use special formatting
                p = doc.add_paragraph()
                # Remove markdown formatting and add as bold
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                run = p.add_run(clean_text)
                run.bold = True
                # Try to apply a specific style if available
                try:
                    p.style = 'Quote'  # or 'Intense Quote' if available
                except:
                    pass  # Style not available, continue with bold formatting
                    
            elif re.match(r'^\*\*(.*?)\*\*', text):
                # Other bold headers - check if they're section headers
                if any(keyword in text.upper() for keyword in ['CALL TO ORDER', 'COMMUNICATIONS', 'CONSENT AGENDA', 'ACTION ITEM', 'DISCUSSION', 'INFORMATION', 'ADJOURN']):
                    # This is a section header
                    header_text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                    p = doc.add_heading(header_text, level=1)
                else:
                    # Regular bold paragraph
                    p = doc.add_paragraph()
                    self.add_formatted_text_to_paragraph(p, text)
                    
            elif text.startswith('>'):
                # Block quote - remove > and make italic
                clean_text = text.lstrip('> ')
                p = doc.add_paragraph()
                run = p.add_run(clean_text)
                run.italic = True
                try:
                    p.style = 'Quote'
                except:
                    pass
                    
            else:
                # Regular paragraph
                p = doc.add_paragraph()
                self.add_formatted_text_to_paragraph(p, text)
                
        except Exception as e:
            # Fallback to simple paragraph
            doc.add_paragraph(text)
    
    def add_formatted_text_to_paragraph(self, paragraph, text):
        """Add text with inline formatting (bold, italic) to a paragraph"""
        try:
            # Split text by formatting markers
            parts = []
            current_pos = 0
            
            # Find all bold (**text**) and italic (*text*) patterns
            bold_pattern = r'\*\*(.*?)\*\*'
            italic_pattern = r'(?<!\*)\*([^*]+?)\*(?!\*)'
            
            # Process bold formatting first
            for match in re.finditer(bold_pattern, text):
                # Add text before bold
                if match.start() > current_pos:
                    parts.append(('normal', text[current_pos:match.start()]))
                # Add bold text
                parts.append(('bold', match.group(1)))
                current_pos = match.end()
            
            # Add remaining text
            if current_pos < len(text):
                remaining_text = text[current_pos:]
                
                # Process italic in remaining text
                italic_parts = []
                current_italic_pos = 0
                
                for match in re.finditer(italic_pattern, remaining_text):
                    # Add text before italic
                    if match.start() > current_italic_pos:
                        italic_parts.append(('normal', remaining_text[current_italic_pos:match.start()]))
                    # Add italic text
                    italic_parts.append(('italic', match.group(1)))
                    current_italic_pos = match.end()
                
                # Add final remaining text
                if current_italic_pos < len(remaining_text):
                    italic_parts.append(('normal', remaining_text[current_italic_pos:]))
                
                parts.extend(italic_parts if italic_parts else [('normal', remaining_text)])
            
            # If no formatting found, add as normal text
            if not parts:
                parts = [('normal', text)]
            
            # Add runs to paragraph
            for format_type, content in parts:
                if content.strip():  # Only add non-empty content
                    run = paragraph.add_run(content)
                    if format_type == 'bold':
                        run.bold = True
                    elif format_type == 'italic':
                        run.italic = True
                        
        except Exception as e:
            # Fallback: add as plain text
            paragraph.add_run(text)
    
    def is_heading(self, line):
        """Determine if a line should be treated as a heading"""
        # Skip markdown-style headings (handled separately)
        if line.startswith('#'):
            return False
            
        # Skip lines that start with formatting markers
        if line.startswith('**') or line.startswith('*'):
            return False
            
        # Headings are typically:
        # - Short (less than 100 characters)
        # - Don't end with punctuation (except :)
        # - May be all caps or title case
        # - May start with numbers (1., I., A., etc.)
        
        if len(line) > 100:
            return False
        
        # Check for common heading patterns
        heading_patterns = [
            r'^\d+\.\s+',  # 1. 2. 3.
            r'^[IVX]+\.\s+',  # I. II. III.
            r'^[A-Z]\.\s+',  # A. B. C.
            r'^[A-Z][A-Z\s]+:?',  # ALL CAPS
            r'^\w+\s*(REPORT|MINUTES|AGENDA|DISCUSSION|ACTION|MOTION)',  # Common meeting terms
            r'^(CALL TO ORDER|COMMUNICATIONS|CONSENT AGENDA|ACTION ITEMS|INFORMATION ITEMS|ADJOURN)',  # Meeting sections
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        
        # If line ends with : and is reasonably short, likely a heading
        if line.endswith(':') and len(line) < 80:
            return True
        
        # If line is short and doesn't end with sentence punctuation
        if len(line) < 60 and not line.endswith(('.', '!', '?')):
            # Check if it looks like a title (most words capitalized)
            words = line.split()
            if len(words) > 1:
                capitalized_words = sum(1 for word in words if word[0].isupper())
                if capitalized_words >= len(words) * 0.7:  # 70% or more words capitalized
                    return True
        
        return False
    
    def get_heading_level(self, line):
        """Determine the heading level (1-3) based on the line content"""
        # Level 1: Main sections, all caps, or numbered with single digit
        if re.match(r'^[1-9]\.\s+', line) or line.isupper():
            return 1
        
        # Level 2: Subsections, lettered, or roman numerals
        if re.match(r'^[A-Z]\.\s+|^[IVX]+\.\s+', line):
            return 3
        
        # Level 3: Everything else
        return 4  # Default to level 1 for simplicity
    
    def generate_minutes_thread(self):
        """Run minutes generation in separate thread"""
        thread = threading.Thread(target=self.generate_minutes)
        thread.daemon = True
        thread.start()
    
    def generate_minutes(self):
        """Main function to generate minutes"""
        try:
            # Validate inputs
            if not self.agenda_path.get():
                self.log("Please select an agenda file.")
                return
            
            if not self.transcript_path.get():
                self.log("Please select a transcript file.")
                return
            
            if not self.output_dir.get():
                self.log("Please select an output directory.")
                return
            
            if not self.api_key_var.get():
                self.log("Please enter your OpenAI API key in settings.")
                return
            
            # Start progress bar
            self.progress.start()
            self.generate_button.config(state='disabled')
            
            self.log("Starting minutes generation...")
            
            # Extract text from files
            self.log("Extracting text from agenda...")
            agenda_text = self.extract_word_text(self.agenda_path.get())
            if not agenda_text:
                self.log("Failed to extract agenda text.")
                return
            
            self.log("Extracting text from transcript...")
            transcript_text = self.extract_vtt_text(self.transcript_path.get())
            if not transcript_text:
                self.log("Failed to extract transcript text.")
                return
            
            # Generate minutes
            minutes_text = self.generate_minutes_chatgpt(agenda_text, transcript_text)
            if not minutes_text:
                return
            
            # Create output filename
            agenda_filename = Path(self.agenda_path.get()).stem
            output_filename = f"{agenda_filename}_minutes.docx"
            output_path = Path(self.output_dir.get()) / output_filename
            
            # Save to Word document
            self.log("Creating Word document...")
            self.minutes_to_word_template(minutes_text, str(output_path))
            
            self.log("Minutes generation completed successfully!")
            messagebox.showinfo("Success", f"Minutes saved to:\n{output_path}")
            
        except Exception as e:
            self.log(f"Unexpected error: {str(e)}")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
        
        finally:
            # Stop progress bar and re-enable button
            self.progress.stop()
            self.generate_button.config(state='normal')

def main():
    root = tk.Tk()
    app = BoardMinutesBuddy(root)
    root.mainloop()

if __name__ == "__main__":
    main()
